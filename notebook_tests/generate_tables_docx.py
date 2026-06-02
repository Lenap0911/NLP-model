"""
Generate appendix finding tables as a Word document.
Run from project root: python generate_tables_docx.py
Output: output/tables/appendix_tables.docx
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(ROOT, 'output', 'tables')
os.makedirs(OUT_DIR, exist_ok=True)

# ── colour helpers ─────────────────────────────────────────────────────────────
HDR_BG  = '1F3A5F'
ROW_A   = 'EEF3F9'
ROW_B   = 'FFFFFF'

def set_cell_bg(cell, hex_colour):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_colour)
    tcPr.append(shd)

def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '9EAEC0')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_table(doc, title, subtitle, col_headers, col_widths_cm, rows):
    # title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_sub = p_sub.add_run(subtitle)
        run_sub.italic = True
        run_sub.font.size = Pt(8.5)
        run_sub.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p_sub.paragraph_format.space_after = Pt(4)

    n_cols = len(col_headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'

    # set column widths
    for i, w in enumerate(col_widths_cm):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    # header row
    hdr_row = table.rows[0]
    for i, (hdr, cell) in enumerate(zip(col_headers, hdr_row.cells)):
        set_cell_bg(cell, HDR_BG)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(hdr)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = ROW_A if r_idx % 2 == 0 else ROW_B
        for c_idx, (text, cell) in enumerate(zip(row_data, row.cells)):
            set_cell_bg(cell, bg)
            set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(str(text))
            run.font.size = Pt(8.5)

    doc.add_paragraph()  # spacer between tables


# ── build document ─────────────────────────────────────────────────────────────
doc = Document()

# page margins: narrow for wide tables
section = doc.sections[0]
section.page_width  = Cm(29.7)   # A4 landscape
section.page_height = Cm(21.0)
section.left_margin   = Cm(1.5)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)

# document title
t = doc.add_heading('Appendix — Pipeline Findings Summary Tables', level=1)
t.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
doc.add_paragraph(
    'Tables A1–A6 summarise the key findings from each stage of the NLP pipeline, '
    'their analytical significance, complementary evidence from the survey and audit, '
    'and supporting literature. Tables follow the pipeline data flow order.'
).runs[0].font.size = Pt(9)
doc.add_paragraph()

COLS4 = ['Finding', 'Analytical significance', 'Complementary evidence', 'Source']
W4    = [6.0, 6.5, 9.0, 3.0]   # cm, total ~24.5 cm usable on A4 landscape

# ── Table A1 ───────────────────────────────────────────────────────────────────
add_table(doc,
    'Table A1 — Corpus Composition and Preprocessing',
    'Pipeline Stage 1: Data collection → preprocessing → language mapping',
    COLS4, W4,
    [
        ['607 articles retained (PT: 313, 52%; ES: 200, 33%; EN: 94, 15%) across 44 flood events in 17 countries.',
         'Working corpus is clean and deduplicated. Portuguese dominates due to Brazil\'s flood frequency; English is the smallest group despite covering North American events.',
         'Audit: 84% structural missingness at collection stage — 226/269 queried events returned zero CDX results. The 607 articles are a heavily filtered subset of potential flood coverage.',
         'Blomeier et al. (2025)'],
        ['North America: 76 articles (12.5%). South America: 531 articles (87.5%).',
         'Regional imbalance shapes all cross-regional comparisons. The regional variable is confounded with language and source type throughout the analysis.',
         'Audit: North America achieves 22.9% crawl coverage vs 14.5% for Latin America at collection stage. The NLP corpus inverts this because of event frequency, not coverage quality.',
         'Barocas et al. (2023)'],
        ['14 post-filter rules applied after Stage 09 in the data collection pipeline.',
         'Filter rules encode assumptions about valid flood articles, conditioning corpus composition before any NLP scoring begins.',
         'Audit: feedback loop — filtering shapes what the model treats as signal. Larger-displacement events are somewhat more likely to be covered (mild signal, not statistically significant).',
         'Barrio Andrés (2021)'],
    ]
)

# ── Table A2 ───────────────────────────────────────────────────────────────────
add_table(doc,
    'Table A2 — Actionability Distribution',
    'Pipeline Stage 2: Sentence-level scoring → article-level aggregation',
    COLS4, W4,
    [
        ['506/607 articles (83.4%) score zero actionability across all sentences.',
         'Near-universal absence of advisory content. The distribution is bimodal: articles either contain no actionable sentences or a detectable minority.',
         'Survey: Article A (high actionability) rated highest for clarity (4.31) and coping confidence (4.08) in isolated evaluation — confirms the pipeline\'s top scorer aligns with human quality judgement.',
         'Lindell & Perry (2012)'],
        ['Corpus mean: 1.99%. Non-zero median: 10.0%. Maximum: 40%.',
         'Even including the non-zero tail, average advisory content is negligible. The mean is partly an artefact of how dominant the advice feature is when present.',
         'Audit: advice feature acts as a near-binary switch at scale. Weights calibrated on a small sample underestimated advice dominance. Pipeline output behaves as a discrete index, not a continuous measure.',
         'Mostafiz et al. (2022)'],
        ['Sentence-level probability distribution: mass near 0.00–0.15, sparse middle (0.15–0.70), small cluster near 1.0.',
         'Structural flaw in feature weights: advice feature compresses most sentences to near-zero and pushes a small subset near one. Score is bimodal not Gaussian.',
         'Audit: two methodological caveats — (1) actionability % is sensitive to article length; (2) text segmentation quality in scraped data skews sentence counts. Recalibration on a larger sample required.',
         'Zguir et al. (2025)'],
        ['Survey ranking paradox: medium-actionability article (Article B) preferred as most useful (mean rank 1.64, 95% CI: 1.53–1.75). High-actionability Article A not preferred despite highest clarity score.',
         'Maximising the pipeline\'s actionability score does not maximise perceived usefulness. Readers prefer narrative context combined with guidance over pure directive density.',
         'NLP pipeline: the non-zero median (10%) sits exactly in the preferred utility range from the survey. Pipeline and survey converge on a threshold, not a maximum.',
         'Palen & Hughes (2018)'],
    ]
)

# ── Table A3 ───────────────────────────────────────────────────────────────────
COLS_A3 = ['Component', 'EN', 'ES', 'PT', 'χ² test', 'Analytical significance', 'Complementary evidence', 'Source']
W_A3    = [3.0, 1.0, 1.0, 1.0, 2.8, 5.0, 6.5, 2.5]

add_table(doc,
    'Table A3 — PADM Component Presence by Language',
    'Pipeline Stage 2: Feature extraction per sentence → article-level aggregation by language group',
    COLS_A3, W_A3,
    [
        ['Spatial anchors', '68%', '88%', '92%', 'χ²=36.19, p<.001',
         'Near-universal in ES/PT. Satisfies PADM\'s named-location criterion but not the directive criterion.',
         'Cluster 1 (actionable, 6%) has highest spatial density (0.599) yet most articles with spatial anchors score zero — location alone is insufficient.',
         'Lindell & Perry (2012)'],
        ['Short-term urgency', '51%', '56%', '56%', 'χ²=0.65, ns',
         'Uniformly distributed across all languages. No cross-linguistic difference.',
         'Urgency co-occurs with spatial in Cluster 0 (90% of corpus, mean 0.7%) without producing actionability — urgency + location are necessary but not sufficient.',
         'Zade et al. (2018)'],
        ['Imperative signals (keyword + POS)', '34%', '10%', '16%', 'χ²=27.67, p<.001',
         'English leads on imperative detection. Spanish and Portuguese lower because conjugated imperative forms differ from infinitive-based lexicons.',
         'Audit: lexicon parity distorted English baseline (KS=0.207). Language-specific scoring logic required; English uses distributed modals, ES/PT use verbal inflections.',
         'Klein (2024)'],
        ['Advice-framing verbs', '7%', '13%', '5%', 'χ²=9.34, p=.009',
         'Rarest PADM component across all languages. Most directly associated with protective action recommendations.',
         'Survey: Article A (high advice-framing) rated clearest in isolation but not preferred for utility. Advice-framing correlates with clarity, not utility preference.',
         'Jurafsky (2014)'],
        ['Cross-linguistic pattern', 'Strong spatial + urgency; near-absent advice', 'Strong spatial + urgency; slightly higher advice', 'Highest spatial; lowest advice', 'Genre convergence',
         'All three languages share the same structural deficit. The failure is a genre-level property of flood journalism, not a language-specific artefact.',
         'Audit: lexicon synchronisation for ES/PT (KS=0.104) still could not produce parity with English — distributional gaps reflect underlying syntactic architecture.',
         'Blomeier et al. (2025)'],
    ]
)

# ── Table A4 ───────────────────────────────────────────────────────────────────
add_table(doc,
    'Table A4 — Source Authority and Regional Analysis',
    'Pipeline Stages 3 & 5: Source authority classification → global region assignment',
    COLS4, W4,
    [
        ['National news mean: 3.61% (n=104); all other outlets: 1.65% (n=503). Mann-Whitney U=30362, p<.001.',
         'Source type is the strongest single predictor of actionability. Institutional national journalism consistently produces more advisory content regardless of language or region.',
         'Survey: high-actionability Article A corresponds to an institutional source context, confirming alignment between source type and human clarity rating.',
         'Semetko & Valkenburg (2000)'],
        ['Source type means: national 3.61% > local 2.49% > regional 2.04% > unknown 1.44% > government agency 0.95% > NGO 0.84% > radio 0.00%.',
         'Source scope (national > regional > local) tracks actionability. Government agencies and NGOs score below average despite their advisory mandate.',
         'Audit: domain lists and keyword thresholds in data collection favour nationally-scoped institutional media — the source type gradient partly reflects corpus construction choices.',
         'Barocas et al. (2023)'],
        ['North America mean: 0.82% (n=76); South America mean: 2.15% (n=531). Mann-Whitney U, p=.147 (not significant).',
         'The regional gap is not statistically significant when tested directly. H₁ is not supported at the regional level.',
         'Audit: North America achieves 22.9% crawl coverage but only 12.5% of the NLP corpus — the North American sample is underpowered for regional comparisons.',
         'He et al. (2024)'],
        ['Regional gap explained by corpus composition: Mexican/Colombian national outlets drive the South American actionability signal; US/Canadian corpus dominated by wire services.',
         'Actionability is a property of institutional national journalism, independent of geography or emergency infrastructure level. H₁ is reframed rather than confirmed or rejected.',
         'Audit: post-filter rules encoding assumptions favouring nationally-scoped media directly shape which source types are over-represented by region.',
         'Quarantelli (2008)'],
    ]
)

# ── Table A5 ───────────────────────────────────────────────────────────────────
add_table(doc,
    'Table A5 — Frame Classification and Source Type Mediation',
    'Pipeline Stage 4: Frame classification → cross-tabulation with source authority',
    COLS4, W4,
    [
        ['Frame distribution: impact 268 (44%), response 169 (28%), accountability 144 (24%), recovery 26 (4%). Kruskal-Wallis H=17.84, p<.001.',
         'Impact framing dominates. Overall frame effect is statistically significant but small (η²=0.025). PADM-relevant response framing is a minority at 28%.',
         'Frame keyword lexicons were not balanced across languages in the audit — classification may carry the same lexicon sensitivity bias as actionability scoring.',
         'Entman (1993)'],
        ['Response framing does not outperform impact framing: 1.97% vs 1.25% (pairwise ns).',
         'Key null result for PADM theory. Response-framed articles describe evacuation orders and emergency resources as events to third parties, not as reader directives.',
         'Cluster 0 (90% of corpus) contains many response-framed articles that are linguistically complete but behaviourally inert. Survey Article B (preferred for utility) was response-framed.',
         'Lindell & Perry (2012)'],
        ['Accountability × national news: 6.81% vs accountability × other outlets: 2.30%. Mann-Whitney U=2768, p<.001, |r|=0.331.',
         'The accountability-actionability correlation is entirely mediated by source type. Accountability in non-national outlets scores no higher than impact or response.',
         'Survey: participants\' utility preferences were independent of framing — their preference for Article B further challenges frame as an independent predictor.',
         'Semetko & Valkenburg (2000)'],
        ['Impact, response, recovery: national vs other differences all not significant (all p>.40).',
         'For three of four frames, source type makes no difference. Only accountability shows the national news premium — frame alone has negligible predictive power once source type is controlled.',
         'Audit: Spanish national accountability journalism (n=37) drives this finding. English and Portuguese national outlets contribute near-zero accountability actionability.',
         'Blomeier et al. (2025)'],
        ['Spanish national accountability: mean 7.34% (n=37). English national: ~0%. Portuguese national (n=3): ~0%.',
         'The corpus-level frame finding reduces to one language group and one source type. Not a generalisable framing effect.',
         'Audit: structural missingness means the English national news accountability sample is effectively zero — cross-linguistic comparison on this finding is not possible from this corpus.',
         'Klein (2024)'],
    ]
)

# ── Table A6 ───────────────────────────────────────────────────────────────────
add_table(doc,
    'Table A6 — Clustering Results',
    'Pipeline Stage 5: Z-score normalisation → K-Means (k=3) on structural features → post-hoc actionability observation',
    COLS4, W4,
    [
        ['K=3 optimal by silhouette score (0.499 vs k=4: 0.286, k=5: 0.208). Structural features only used; actionability excluded from clustering input.',
         'Three distinct article profiles emerge without actionability as input — structure is real, not circular. Kruskal-Wallis H=275.5, p<.001, η²=0.453 — largest effect size in the entire analysis.',
         'All three cluster pairs differ significantly (p<.001 pairwise). η²=0.453 is large — cluster membership explains 45% of variance in actionability scores.',
         'Sit et al. (2020)'],
        ['Cluster 0 (Descriptive Baseline): n=549, 90% of corpus, mean actionability 0.7%. Near-zero on all five PADM structural features.',
         'Nine in ten articles belong to a cluster with no advisory features. Flood journalism\'s default state is factually complete and behaviourally inert.',
         'Survey: Article C (low actionability, Cluster 0 profile) scored lowest on clarity (3.12) and coping confidence (3.04). Pipeline classification aligns with human evaluation in isolation.',
         'Zade et al. (2018)'],
        ['Cluster 1 (Actionable Advisory): n=39, 6% of corpus, mean actionability 18.0%. Defined by advice-framing (0.185) and spatial density (0.599).',
         'Only cluster approaching PADM compliance. Driven by advice-framing verbs — not imperatives or urgency, both of which are present in other clusters without comparable actionability.',
         'Survey ranking paradox: even Cluster 1 articles (~18%) may not be the most useful — survey participants preferred medium-actionability (~10%). Cluster 1 exceeds the utility threshold.',
         'Lindell & Perry (2012)'],
        ['Cluster 2 (Recovery Discourse): n=19, 3% of corpus, mean actionability 5.3%. Defined by long-term recovery keywords (0.456).',
         'Recovery-oriented articles form a distinct profile but score well below Cluster 1. Forward-oriented institutional language does not satisfy PADM\'s immediate protective action criterion.',
         'Recovery is the smallest frame category (n=26, 4%) and the cluster is also the smallest. Findings should be interpreted cautiously due to low n.',
         'Mostafiz et al. (2022)'],
        ['Advice-framing defines Cluster 1 but is absent in 87–95% of articles across all three languages.',
         'The sole PADM-compliant article type is rare, linguistically specific, and concentrated in Spanish national news. It is not a genre-level property but an outlier driven by institutional source type.',
         'Audit: advice detection is harder in English (modal-based directives) than in Spanish/PT (inflection-based). Cluster 1\'s Spanish dominance may partly reflect detection sensitivity differences.',
         'Klein (2024)'],
    ]
)

# ── save ───────────────────────────────────────────────────────────────────────
out_path = os.path.join(OUT_DIR, 'appendix_tables.docx')
doc.save(out_path)
print(f'saved: {out_path}')
